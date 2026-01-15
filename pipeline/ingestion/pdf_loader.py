from langchain_community.document_loaders import PyPDFLoader, TextLoader
from pathlib import Path
import docx
import csv
import xml.etree.ElementTree as ET

def process_all_docs(upload_dir: str):
    all_documents = []
    base_dir = Path(upload_dir)

    # --- PDFs ---
    pdf_files = list(base_dir.glob("**/*.pdf"))
    for pdf_file in pdf_files:
        loader = PyPDFLoader(str(pdf_file))
        docs = loader.load()
        for doc in docs:
            doc.metadata["source_file"] = pdf_file.name
            doc.metadata["file_type"] = "pdf"
        all_documents.extend(docs)

    # --- TXT files ---
    txt_files = list(base_dir.glob("**/*.txt"))
    for txt_file in txt_files:
        loader = TextLoader(str(txt_file))
        docs = loader.load()
        for doc in docs:
            doc.metadata["source_file"] = txt_file.name
            doc.metadata["file_type"] = "txt"
        all_documents.extend(docs)

    # --- DOCX files ---
    docx_files = list(base_dir.glob("**/*.docx"))
    for docx_file in docx_files:
        doc_content = ""
        doc_obj = docx.Document(str(docx_file))
        for para in doc_obj.paragraphs:
            doc_content += para.text + "\n"

        # Wrap into a LangChain-style Document
        from langchain.schema import Document
        doc = Document(page_content=doc_content, metadata={
            "source_file": docx_file.name,
            "file_type": "docx"
        })
        all_documents.append(doc)

    # --- CSV files ---
    csv_files = list(base_dir.glob("**/*.csv"))
    for csv_file in csv_files:
        doc_content = ""
        with open(csv_file, newline="", encoding="utf-8") as f:
            reader = csv.reader(f)
            for row in reader:
                doc_content += " | ".join(row) + "\n"
        from langchain.schema import Document
        doc = Document(page_content=doc_content, metadata={
            "source_file": csv_file.name,
            "file_type": "csv"
        })
        all_documents.append(doc)

    # --- XML files ---
    xml_files = list(base_dir.glob("**/*.xml"))
    for xml_file in xml_files:
        tree = ET.parse(xml_file)
        root = tree.getroot()
        doc_content = ET.tostring(root, encoding='unicode')
        from langchain.schema import Document
        doc = Document(page_content=doc_content, metadata={
            "source_file": xml_file.name,
            "file_type": "xml"
        })
        all_documents.append(doc)

    print(f"Total documents loaded: {len(all_documents)}")
    return all_documents
